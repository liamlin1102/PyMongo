#-*- coding: UTF-8 -*-

from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
from sklearn.model_selection import train_test_split

import datetime
from functools import reduce
import re
import textwrap
import seaborn as seabornInstance
import matplotlib.pyplot as plt

import statsmodels.api as sm
from statsmodels.iolib import summary,table
from statsmodels.iolib.table import SimpleTable
from statsmodels.iolib.tableformatting import fmt_latex, fmt_txt
from statsmodels.iolib.summary2 import _make_unique,_col_info,summary_params,Summary
from statsmodels.compat.python import lzip

import docx
from docx.shared import Inches,Cm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import sys, dateutil.parser, time, math, requests, re, json, os
import pandas as pd
import numpy as np
from datetime import datetime as dt
from collections import defaultdict
from bson import json_util
from bson.json_util import loads

#資料處理轉換
class NpEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.integer):
            return int(obj)
        elif isinstance(obj, np.floating):
            return float(obj)
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        else:
            return super(NpEncoder, self).default(obj)

def stdtime(daytime):
    get=dateutil.parser.parse(daytime)
    return get

# 2021.3.17 Lin修改interval,intervals
def interval(start_time,end_time,freq):
    data = pd.date_range(start=start_time,end=end_time,freq=freq)
    data = pd.DataFrame(data,columns=['stdtime'])
    data[freq]=1
    return data

def intervals(start_time,end_time,freq_list):
    data = pd.date_range(start=start_time,end=end_time,freq="1s")
    data = pd.DataFrame(data,columns=['stdtime'])
    for sec in freq_list:
        final = data[["stdtime"]].set_index("stdtime").resample(sec,origin="start").ffill().reset_index().rename(columns={"index":"stdtime"})
        final[sec]=1
        data = pd.merge(data,final,on="stdtime",how="left")
        data[sec] = data[sec].fillna(0).astype(int)
    return data

def lag(title,phase,data):
    for p in phase:
        df=data[title].shift(p)
        for t in title:
            df.rename(columns={t:"lag"+str(p)+"_"+t},inplace=True)
        data = pd.concat([data,df],axis=1)
    data = data.dropna().reset_index(drop=True)
    return data

#計算敘述統計
def status(data):
    data = pd.DataFrame([data.sum(),data.count(),data.min(),data.idxmin(),data.quantile(.25),data.median(),
                      data.quantile(.75),data.mean(),data.max(),data.idxmax(),data.mad(),data.var(),
                      data.std(),data.skew(),data.kurt()],index=['加總','總数','最小值','最小值位置','25%分位数',
                    '中位数','75%分位数','均值','最大值','最大值位数','平均絕對偏差','方差','標準差','偏度','峰度'])
    return data


# 21.03.30 Cao 新增 OLS回归def
# 回归OLS部分 _col_params()和summary_col()函数共同使用

# %f ——保留小数点后面六位有效数字，%.3f，保留3位小数位
# %e ——保留小数点后面六位有效数字，指数形式输出，%.3e，保留3位小数位，使用科学计数法
# %g ——在保证六位有效数字的前提下，使用小数方式，否则使用科学计数法，%.3g，保留3位有效数字，使用小数或科学计数法
def _col_params(result, 
                float_format_t='%.4f',
                float_format_coef='%.2f',
                float_format_adjr='%.2f',
                stars=True):

    # Extract parameters
    res = summary_params(result)
    # Format float
    res['Coef.'] = res['Coef.'].apply(lambda x: float_format_coef % x)
    res['t'] = res['t'].apply(lambda x: float_format_t % x)
#     for col in res.columns[:3]:
#         res[col] = res[col].apply(lambda x: float_format % x)
    # Std.Errors in parentheses
    res.iloc[:, 2] = '(' + res.iloc[:, 2] + ')'
    # Significance stars
    if stars:
        idx = res.iloc[:, 3] < .1
        res.loc[idx, res.columns[0]] = res.loc[idx, res.columns[0]] + '*'
        idx = res.iloc[:, 3] < .05
        res.loc[idx, res.columns[0]] = res.loc[idx, res.columns[0]] + '*'
        idx = res.iloc[:, 3] < .01
        res.loc[idx, res.columns[0]] = res.loc[idx, res.columns[0]] + '*'
    # Stack Coefs and Std.Errors
    res = res.iloc[:, [0,2]]
    res = res.stack()
    rsquared = getattr(result, 'rsquared', np.nan)
    rsquared_adj = getattr(result, 'rsquared_adj', np.nan)
    r2 = pd.Series({('R-squared', ""): rsquared,
                    ('R-squared Adj.', ""): rsquared_adj})

    if r2.notnull().any():
        r2 = r2.apply(lambda x: float_format_adjr % x)
        res = pd.concat([res, r2], axis=0)
    res = pd.DataFrame(res)
    res.columns = [str(result.model.endog_names)]
    return res



def summary_col(results,
                float_format_t='%.4f',
                float_format_coef='%.2f',
                float_format_adjr='%.2f',
                model_names=(),
                stars=False,
                info_dict=None, regressor_order=(), drop_omitted=False):
    # sample
    # y = df['nft']
    # x = df[['epm',"volt","lnvol","relative_spread"]] 
    # x1 = sm.add_constant(x)
    # est = sm.OLS(y, x1).fit()
    # a = summary_col([est,est],stars=True,float_format_t='%0.2f',float_format_coef='%0.4f',float_format_adjr='%0.2f')

    
    if not isinstance(results, list):
        results = [results]

    cols = [_col_params(x, stars=stars, float_format_t=float_format_t,float_format_coef=float_format_coef,float_format_adjr=float_format_adjr) for x in
            results]

    # Unique column names (pandas has problems merging otherwise)
    if model_names:
        colnames = _make_unique(model_names)
    else:
        colnames = _make_unique([x.columns[0] for x in cols])
    for i in range(len(cols)):
        cols[i].columns = [colnames[i]]

    def merg(x, y):
        return x.merge(y, how='outer', right_index=True,
                       left_index=True)

    summ = reduce(merg, cols)

    if regressor_order:
        varnames = summ.index.get_level_values(0).tolist()
        vc = pd.Series(varnames).value_counts()
        varnames = vc.loc[vc == 2].index.tolist()
        ordered = [x for x in regressor_order if x in varnames]
        unordered = [x for x in varnames if x not in regressor_order]
        new_order = ordered + unordered
        other = [x for x in summ.index.get_level_values(0)
                 if x not in new_order]
        new_order += other
        if drop_omitted:
            for uo in unordered:
                new_order.remove(uo)
        summ = summ.loc[new_order]

    idx = []
    index = summ.index.get_level_values(0)
    for i in range(0, index.shape[0], 2):
        idx.append(index[i])
        if (i + 1) < index.shape[0] and (index[i] == index[i + 1]):
            idx.append("")
        else:
            idx.append(index[i + 1])
    summ.index = idx

    # add infos about the models.
    if info_dict:
        cols = [_col_info(x, info_dict.get(x.model.__class__.__name__,
                                           info_dict)) for x in results]
    else:
        cols = [_col_info(x, getattr(x, "default_model_infos", None)) for x in
                results]
    # use unique column names, otherwise the merge will not succeed
    for df, name in zip(cols, _make_unique([df.columns[0] for df in cols])):
        df.columns = [name]

    def merg(x, y):
        return x.merge(y, how='outer', right_index=True,
                       left_index=True)

    info = reduce(merg, cols)
    dat = pd.DataFrame(np.vstack([summ, info]))  # pd.concat better, but error
    dat.columns = summ.columns
    dat.index = pd.Index(summ.index.tolist() + info.index.tolist())
    
    dat =dat.drop(index=["R-squared"])
    dat = dat.reset_index()
    return dat

# 输出回归为docx格式
def output_docx(dataframe,output_name):
    # 表格样式固定为"Light Shading" 样式详见 http://www.voidcn.com/article/p-weenhbxd-bqy.html
    # Sample
    # output_docx(a,path+"a.docx")
    doc = docx.Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman' #  字体
    font.size = Pt(14)  # 文字大小

    # 按照行列添加表格
    t = doc.add_table(dataframe.shape[0]+1, dataframe.shape[1],style="Light Shading")
    
    # 表头
    for j in range(dataframe.shape[-1]):
        t.cell(0,j).text = dataframe.columns[j]
        t.cell(0,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  #居中对齐
        t.cell(0,j).paragraphs[0].paragraph_format.space_before=Pt(4)  #上行间距
        t.cell(0,j).paragraphs[0].paragraph_format.space_after=Pt(4)   #下行间距

    # 表身
    for i in range(dataframe.shape[0]):
        for j in range(dataframe.shape[-1]):
            t.cell(i+1,j).text = str(dataframe.values[i,j])
            t.cell(i+1,j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(i+1,j).paragraphs[0].paragraph_format.space_before=Pt(4)
            t.cell(i+1,j).paragraphs[0].paragraph_format.space_after=Pt(4)

    # 表格行高
    t.rows[0].height=Pt(30)
    for i in range(1,len(a)+1):
        t.rows[i].height=Pt(20) 
    
    p = doc.add_paragraph()
    run=p.add_run('t-value in parentheses.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(20)
    p.paragraph_format.space_after=Pt(0)
    
    p = doc.add_paragraph()
    run=p.add_run('* 1.65<t<1.96, ** 1.96<t<2.33, *** 2.33<t')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)  
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    p.paragraph_format.space_before=Pt(0)
    p.paragraph_format.space_after=Pt(0)
    
    # 输出
    doc.save(output_name)